# ruff: noqa: E501
from __future__ import annotations

import json
from collections.abc import Iterable, Sequence
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_DIR / "docs" / "Bao_Cao_FaceTrust_Hoan_Thien.docx"
ASSET_DIR = PROJECT_DIR / "docs" / "report-assets"
CROSS_RESULTS = PROJECT_DIR / "reports" / "cross_dataset_results.json"
PIPELINE_RESULTS = PROJECT_DIR / "reports" / "evaluation_results.json"
TOC_MAP_PATH = PROJECT_DIR / "docs" / "report_toc_pages.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "102136"
MUTED = "5E6C7E"
TEAL = "0D8A8F"
GREEN = "15805D"
RED = "9B1C1C"
GOLD = "7A5A00"
LIGHT_GRAY = "F4F6F9"
PALE_BLUE = "E8F2F6"
PALE_TEAL = "E7F5F3"
PALE_GOLD = "FFF4DD"
LINE = "C9D5E1"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], *, indent: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_rule(paragraph, color: str = TEAL, size: int = 8) -> None:
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(p_bdr)


def paragraph_shading(paragraph, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, display, end))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    for key in ("ascii", "hAnsi", "eastAsia"):
        normal._element.rPr.rFonts.set(qn(f"w:{key}"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        for key in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{key}"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("FACETRUST  |  BÁO CÁO MÔN XỬ LÝ ẢNH")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    paragraph_rule(header, LINE, 4)
    add_page_number(section.footer.paragraphs[0])


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_run_font(paragraph.add_run(item), color=INK)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    p_pr.extend((tabs, indent))
    level.extend((start, num_fmt, level_text, justification, p_pr))
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num_id_node))
        paragraph._p.get_or_add_pPr().append(num_pr)
        set_run_font(paragraph.add_run(item), color=INK)


def add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = PALE_GOLD,
    border: str = GOLD,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph_shading(paragraph, fill, border)
    set_run_font(paragraph.add_run(f"{label}: "), bold=True, color=border)
    set_run_font(paragraph.add_run(text), color=INK)


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_shading(paragraph, LIGHT_GRAY, TEAL)
    set_run_font(paragraph.add_run(text), name="Consolas", size=9, color=INK)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths: Sequence[int],
    *,
    numeric: set[int] | None = None,
) -> None:
    numeric = numeric or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(value), size=9.2, bold=True, color=INK)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index in numeric else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            set_run_font(paragraph.add_run(str(value)), size=9.2, color=INK)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.35) -> None:
    path = ASSET_DIR / filename
    if not path.exists():
        add_callout(doc, "Thiếu hình", f"Không tìm thấy {path}", fill="FBECE9", border=RED)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    inline = paragraph.add_run().add_picture(str(path), width=Inches(width))
    inline._inline.docPr.set("descr", caption)
    doc.add_paragraph(caption, style="Caption")


def add_chapter(doc: Document, number: int, title: str) -> None:
    doc.add_page_break()
    paragraph = doc.add_paragraph(f"CHƯƠNG {number}. {title.upper()}", style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph_rule(paragraph)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    node = OxmlElement("w:t")
    node.text = text
    run.extend((r_pr, node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def build_cover(doc: Document) -> None:
    for text, size, after, bold, color in (
        ("[TÊN TRƯỜNG]", 12, 3, True, MUTED),
        ("[TÊN KHOA / BỘ MÔN]", 11, 44, False, MUTED),
        ("BÁO CÁO ĐỒ ÁN MÔN XỬ LÝ ẢNH", 12, 12, True, TEAL),
        ("FACETRUST", 32, 10, True, INK),
        ("HỆ THỐNG KIỂM ĐỊNH ẢNH KHUÔN MẶT\nVÀ PHÁT HIỆN DEEPFAKE", 19, 28, True, DARK_BLUE),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        set_run_font(paragraph.add_run(text), size=size, color=color, bold=bold)
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_rule(rule, TEAL, 10)
    rule.paragraph_format.space_after = Pt(42)
    for text in (
        "Nhóm thực hiện: [Điền họ tên - MSSV của từng thành viên]",
        "Lớp: [Điền lớp học phần]",
        "Giảng viên hướng dẫn: [Điền họ tên giảng viên]",
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(7)
        set_run_font(paragraph.add_run(text), size=11, color=INK)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(70)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        paragraph.add_run("Thành phố Hồ Chí Minh, 2026"),
        size=10.5,
        color=MUTED,
        italic=True,
    )


def build_front_matter(doc: Document, cross: dict) -> None:
    doc.add_page_break()
    doc.add_paragraph("THÔNG TIN VÀ CAM KẾT BÁO CÁO", style="Heading 1")
    add_callout(
        doc,
        "Cần hoàn thiện trước khi nộp",
        "Điền tên trường, khoa, lớp, giảng viên, họ tên và MSSV ở trang bìa. Các trường này để trống vì tài liệu nguồn không có thông tin đáng tin cậy.",
    )
    doc.add_paragraph("Phạm vi trách nhiệm", style="Heading 2")
    add_body(
        doc,
        "Báo cáo mô tả đúng phiên bản FaceTrust trong repository tại thời điểm hoàn thiện. Nhóm chịu trách nhiệm chạy lại phép thử, lưu kết quả và trình bày trung thực giới hạn. Số liệu do tác giả checkpoint công bố được ghi rõ là tham khảo, không được nhận là kết quả của nhóm.",
    )
    doc.add_paragraph("Cam kết học thuật", style="Heading 2")
    add_body(
        doc,
        "Nhóm không dùng bộ demo tuyển chọn để công bố độ chính xác tổng quát; không diễn giải raw score hay decision margin thành xác suất chắc chắn đúng; không che giấu trường hợp dự đoán sai. Mọi bảng đánh giá của nhóm đều có script, dữ liệu đầu vào và JSON để kiểm tra.",
    )
    doc.add_paragraph("Phân công nhóm", style="Heading 2")
    add_table(
        doc,
        ["STT", "Họ tên - MSSV", "Nhiệm vụ", "Tỷ lệ"],
        [
            ["1", "[Điền thông tin]", "Dữ liệu, model và benchmark", "[ ]%"],
            ["2", "[Điền thông tin]", "Backend, API và kiểm thử", "[ ]%"],
            ["3", "[Điền thông tin]", "Frontend, tài liệu và trình bày", "[ ]%"],
        ],
        [700, 2300, 4360, 2000],
        numeric={0, 3},
    )

    doc.add_page_break()
    doc.add_paragraph("TÓM TẮT", style="Heading 1")
    add_body(
        doc,
        "FaceTrust là ứng dụng web kiểm định ảnh có khuôn mặt theo ba trạng thái real, fake và uncertain. Hệ thống nhận ảnh, kiểm tra định dạng, định vị khuôn mặt, chuẩn hóa crop, chạy hai detector cục bộ và ánh xạ đầu ra sang kết luận có tín hiệu giải thích. Mục tiêu không phải xác thực danh tính hay đưa ra kết luận pháp lý, mà xây dựng một full pipeline xử lý ảnh có thể tái chạy và đánh giá.",
    )
    add_body(
        doc,
        "Detector chính là checkpoint MS-EffGCViT B0 đã huấn luyện trên FaceForensics++; detector phụ là EfficientNet-B0 dùng có điều kiện khi khuôn mặt chiếm dưới 12% diện tích ảnh. Cả hai trọng số lưu cục bộ trong models. Nhóm không tuyên bố huấn luyện model chính từ đầu; đóng góp nằm ở tích hợp, tiền xử lý, quy tắc quyết định thích nghi, giao diện, kiểm thử và benchmark chéo.",
    )
    overall = cross.get("overall", {})
    if overall:
        add_callout(
            doc,
            "Kết quả chính",
            f"Benchmark chéo trên {overall['total']} frame độc lập với bộ demo đạt strict accuracy {pct(overall['strict_accuracy'])}, fake recall {pct(overall['fake_recall'])}, real recall {pct(overall['real_recall'])} và balanced accuracy {pct(overall['balanced_accuracy'])}. Kết quả cho thấy model có khả năng phát hiện đáng kể nhưng còn thiên lệch và chưa đạt mức sản phẩm pháp chứng.",
            fill=PALE_TEAL,
            border=GREEN,
        )
    add_body(
        doc,
        "Báo cáo tách ba lớp bằng chứng: số liệu tác giả checkpoint, cross-dataset do nhóm chạy và full-pipeline evaluation của API. Cách tách này trả lời câu hỏi kết quả có thật hay không bằng dữ liệu, mã đánh giá, ma trận nhầm lẫn và log từng ảnh thay vì cảm tính.",
    )
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run("Từ khóa: "), bold=True, color=INK)
    set_run_font(
        paragraph.add_run(
            "deepfake detection, face-swap, image forensics, cross-dataset evaluation, FastAPI, EfficientNet, vision transformer."
        ),
        italic=True,
        color=MUTED,
    )


def build_toc(doc: Document) -> None:
    pages = load_json(TOC_MAP_PATH)
    doc.add_page_break()
    doc.add_paragraph("MỤC LỤC", style="Heading 1")
    items = [
        ("CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", "CHƯƠNG 1."),
        ("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", "CHƯƠNG 2."),
        ("CHƯƠNG 3. DỮ LIỆU VÀ NGUỒN GỐC MODEL", "CHƯƠNG 3."),
        ("CHƯƠNG 4. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "CHƯƠNG 4."),
        ("CHƯƠNG 5. TRIỂN KHAI HỆ THỐNG", "CHƯƠNG 5."),
        ("CHƯƠNG 6. EVALUATION, BENCHMARK VÀ KẾT QUẢ", "CHƯƠNG 6."),
        ("CHƯƠNG 7. QUY TRÌNH LÀM VIỆC VỚI AI", "CHƯƠNG 7."),
        ("CHƯƠNG 8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "CHƯƠNG 8."),
        ("PHỤ LỤC A. HƯỚNG DẪN TÁI LẬP", "PHỤ LỤC A."),
        ("PHỤ LỤC B. CÂU HỎI PHẢN BIỆN", "PHỤ LỤC B."),
        ("PHỤ LỤC C. CÔNG THỨC VÀ VÍ DỤ", "PHỤ LỤC C."),
        ("TÀI LIỆU THAM KHẢO", "TÀI LIỆU THAM KHẢO"),
    ]
    add_table(
        doc,
        ["Nội dung", "Trang"],
        [[title, pages.get(key, "-")] for title, key in items],
        [8160, 1200],
        numeric={1},
    )
    doc.add_paragraph("DANH MỤC HÌNH", style="Heading 2")
    add_bullets(
        doc,
        [
            "Hình 3.1. Quy mô tập đánh giá và mẫu ảnh demo",
            "Hình 4.1. Kiến trúc triển khai FaceTrust",
            "Hình 4.2. Full pipeline từ ảnh đầu vào đến kết luận",
            "Hình 4.3. Ánh xạ raw score sang chỉ số rủi ro",
            "Hình 5.1. Giao diện kiểm định ảnh",
            "Hình 6.1. Kết quả cross-dataset",
            "Hình 6.2. Ma trận nhầm lẫn gộp",
            "Hình 7.1. Quy trình làm việc có AI hỗ trợ",
        ],
    )
    doc.add_page_break()
    doc.add_paragraph("DANH MỤC BẢNG", style="Heading 2")
    add_bullets(
        doc,
        [
            "Bảng 2.1. Các dạng thao túng khuôn mặt",
            "Bảng 3.1. Dữ liệu lưu trong dự án",
            "Bảng 3.2. Nguồn gốc và vai trò checkpoint",
            "Bảng 4.1. Các mode lỗi và cách xử lý",
            "Bảng 5.1. Công nghệ triển khai",
            "Bảng 6.1. Ba lớp đánh giá",
            "Bảng 6.2. Kết quả cross-dataset",
            "Bảng 6.3. Ma trận nhầm lẫn",
        ],
    )
    doc.add_paragraph("TỪ VIẾT TẮT", style="Heading 2")
    add_table(
        doc,
        ["Ký hiệu", "Diễn giải"],
        [
            ["API", "Application Programming Interface"],
            ["AUC", "Area Under the ROC Curve"],
            ["CNN", "Convolutional Neural Network"],
            ["FF++", "FaceForensics++"],
            ["FPR", "False Positive Rate"],
            ["GCViT", "Global Context Vision Transformer"],
            ["OOD", "Out-of-distribution"],
            ["TP/FP/TN/FN", "True/False Positive và True/False Negative"],
        ],
        [1800, 7560],
    )


def chapter_1(doc: Document) -> None:
    add_chapter(doc, 1, "Tổng quan đề tài")
    doc.add_paragraph("1.1. Bối cảnh", style="Heading 2")
    add_body(
        doc,
        "Sự phát triển của mô hình sinh ảnh và kỹ thuật hoán đổi khuôn mặt làm nội dung giả ngày càng khó nhận biết bằng mắt. Deepfake có thể phục vụ giải trí nhưng cũng gây nhầm lẫn danh tính, lừa đảo và suy giảm niềm tin đối với bằng chứng hình ảnh. Phát hiện deepfake vì vậy là bài toán phân loại có yêu cầu cao về khả năng khái quát hóa ngoài dữ liệu huấn luyện.",
    )
    add_body(
        doc,
        "Khó khăn cốt lõi là mỗi công cụ tạo giả để lại dấu vết khác nhau. Dấu vết còn thay đổi theo độ phân giải, tỷ lệ khuôn mặt, nén JPEG, chụp lại màn hình và chỉnh sửa sau tạo. Một model đạt điểm rất cao trên tập cùng miền vẫn có thể giảm mạnh với nguồn khác. Đồ án vì vậy tập trung cả vào pipeline dự đoán lẫn cách chứng minh model chạy thật và đánh giá đúng phương pháp.",
    )
    doc.add_paragraph("1.2. Bài toán", style="Heading 2")
    add_body(
        doc,
        "Đầu vào là ảnh JPG, JFIF, PNG hoặc WebP có thể chứa khuôn mặt. Đầu ra gồm nhãn fake, real hoặc uncertain; fake risk index; decision margin; trạng thái phát hiện khuôn mặt; mã phiên quét và tối đa năm tín hiệu giải thích. Hệ thống xử lý ảnh đơn, không xác minh người trong ảnh, không kiểm tra nguồn tin và không thay thế chuyên gia pháp chứng.",
    )
    add_callout(
        doc,
        "Câu hỏi nghiên cứu",
        "Một checkpoint đã huấn luyện trên dữ liệu deepfake có thể được tích hợp thành pipeline web tái lập như thế nào, và khả năng khái quát hóa trên hai nguồn dữ liệu khác miền đạt mức nào?",
        fill=PALE_BLUE,
        border=BLUE,
    )
    doc.add_paragraph("1.3. Mục tiêu", style="Heading 2")
    add_bullets(
        doc,
        [
            "Xây dựng web app tải ảnh và nhận kết luận có giải thích.",
            "Chạy checkpoint AI cục bộ, không trả kết quả theo tên tệp hay hard-code.",
            "Tách bộ demo khỏi benchmark và lưu nguồn ảnh rõ ràng.",
            "Định nghĩa công thức score, ngưỡng và quy tắc kết luận có thể kiểm tra trong mã.",
            "Đánh giá cross-dataset bằng accuracy, balanced accuracy, precision, recall, F1 và ma trận nhầm lẫn.",
            "Cung cấp script tái chạy, JSON từng ảnh, kiểm thử backend và tài liệu phản biện.",
        ],
    )
    doc.add_paragraph("1.4. Phạm vi và giới hạn", style="Heading 2")
    add_body(
        doc,
        "Đồ án tập trung vào ảnh có khuôn mặt và thao túng dạng face-swap/face-forgery. Phiên bản hiện tại chưa tổng hợp bằng chứng theo nhiều frame video, chưa dùng metadata nguồn, chưa có Grad-CAM được kiểm định và chưa hiệu chuẩn xác suất trên calibration set đủ lớn. Kết quả chỉ nên xem là tín hiệu hỗ trợ sàng lọc.",
    )
    add_body(
        doc,
        "Dữ liệu trong repository là frame trích sẵn, không phải toàn bộ video FaceForensics++ hay Celeb-DF-v2. Các frame cùng video có thể tương quan. Báo cáo nêu giới hạn này để tránh biến số lượng frame thành ấn tượng sai về số tình huống độc lập.",
    )
    doc.add_paragraph("1.5. Đóng góp thực tế", style="Heading 2")
    add_numbered(
        doc,
        [
            "Chuyển hệ thống từ hướng bảo vệ ảnh sang nhận diện deepfake theo yêu cầu mới.",
            "Tích hợp MS-EffGCViT B0 và EfficientNet-B0 theo quy tắc khuôn mặt nhỏ.",
            "Thiết kế risk index và decision margin có công thức công khai.",
            "Xây dựng FastAPI và giao diện một trang upload - quét - kết luận.",
            "Tạo đánh giá demo/full pipeline và cross-dataset độc lập, xuất Markdown/JSON.",
            "Tài liệu hóa quy trình AI, giới hạn model và câu trả lời phản biện.",
        ],
    )


def chapter_2(doc: Document) -> None:
    add_chapter(doc, 2, "Cơ sở lý thuyết")
    doc.add_paragraph("2.1. Deepfake và các dạng thao túng", style="Heading 2")
    add_body(
        doc,
        "Deepfake trong phạm vi đồ án là ảnh hoặc frame mà vùng khuôn mặt đã bị tổng hợp hay biến đổi bằng thuật toán. Face-swap thay danh tính nguồn vào khuôn mặt mục tiêu; reenactment truyền biểu cảm; neural rendering tổng hợp texture; ảnh sinh toàn phần tạo cả con người lẫn bối cảnh. Detector hiện tại chuyên face-forgery nên không nên suy rộng kết quả sang mọi loại ảnh AI.",
    )
    add_table(
        doc,
        ["Dạng", "Mô tả", "Dấu vết thường gặp"],
        [
            ["Face-swap", "Thay khuôn mặt người A vào người B", "Biên ghép, màu da, texture và hình học"],
            ["Reenactment", "Truyền biểu cảm hoặc chuyển động", "Vùng miệng, mắt và temporal artifact"],
            ["Neural rendering", "Sinh lại vùng mặt bằng mạng học sâu", "Texture, tần số và nhiễu bất thường"],
            ["Ảnh sinh toàn phần", "Sinh toàn bộ chân dung", "Có thể không có dấu vết face-swap đã học"],
        ],
        [1800, 3400, 4160],
    )
    doc.add_paragraph("2.2. Phân loại nhị phân và uncertain", style="Heading 2")
    add_body(
        doc,
        "Model nền tạo raw score liên tục, sau đó ngưỡng chuyển score thành nhãn. Full pipeline còn yêu cầu phát hiện được khuôn mặt. Nếu không có crop đủ rõ, FaceTrust trả uncertain thay vì ép đoán. Trong đánh giá strict, uncertain tính là sai vì hệ thống chưa đưa đúng ground truth; tỷ lệ uncertain vẫn được báo cáo riêng.",
    )
    doc.add_paragraph("2.3. EfficientNet và GCViT", style="Heading 2")
    add_body(
        doc,
        "EfficientNet dùng compound scaling để cân bằng chiều sâu, chiều rộng và độ phân giải. Bản B0 có chi phí thấp, phù hợp CPU. Checkpoint chính kết hợp EfficientNet với Global Context Vision Transformer để khai thác texture cục bộ và ngữ cảnh toàn cục. Model card ghi đầu vào face-cropped 224×224 và khoảng 8,7 triệu tham số.",
    )
    add_body(
        doc,
        "FaceTrust không đưa toàn khung thẳng vào model chính. Pipeline tìm mặt, mở rộng crop, resize và chuẩn hóa theo mean/std ImageNet. Cách này bám sát model card và giảm việc nền, tóc hoặc quần áo chi phối khi khuôn mặt nhỏ.",
    )
    doc.add_paragraph("2.4. Domain shift và khái quát hóa", style="Heading 2")
    add_body(
        doc,
        "Same-dataset evaluation đo trên phân phối gần tập train; cross-dataset evaluation đo trên nguồn khác. Khác biệt về danh tính, công cụ tạo giả, codec, crop và độ phân giải tạo domain shift. Accuracy 98% trên FF++ không đồng nghĩa 98% với ảnh Internet. Báo cáo ghi nguồn từng con số để không đánh tráo hai loại đánh giá.",
    )
    doc.add_paragraph("2.5. Các chỉ số", style="Heading 2")
    add_code_block(
        doc,
        "Accuracy = (TP + TN) / (TP + TN + FP + FN)\n"
        "Precision_fake = TP / (TP + FP)\n"
        "Recall_fake = TP / (TP + FN)\n"
        "F1_fake = 2 × Precision × Recall / (Precision + Recall)\n"
        "Balanced Accuracy = (Recall_fake + Recall_real) / 2",
    )
    add_body(
        doc,
        "Fake recall trả lời trong toàn bộ ảnh fake, hệ thống bắt được bao nhiêu. Real recall trả lời trong toàn bộ ảnh thật, hệ thống giữ đúng bao nhiêu. Nếu gắn fake cho hầu hết ảnh, fake recall tăng nhưng false positive tăng và real recall giảm. Balanced accuracy giúp nhìn cân bằng hai lớp.",
    )
    add_callout(
        doc,
        "Không nhầm lẫn",
        "Decision margin là khoảng cách tương đối đến ngưỡng đang kích hoạt, không phải xác suất dự đoán đúng. Accuracy là chỉ số của cả tập đã gắn nhãn, không thể suy ra từ một ảnh đơn.",
        fill="FBECE9",
        border=RED,
    )
    doc.add_paragraph("2.6. Calibration và độ tin cậy", style="Heading 2")
    add_body(
        doc,
        "Một sigmoid score không tự động là xác suất đã hiệu chuẩn. Calibration cần một tập giữ riêng, sau đó so sánh score với tần suất đúng thực tế bằng reliability diagram, Expected Calibration Error hoặc Brier score. FaceTrust hiện chưa có bước này nên chỉ công bố risk index và decision margin, đồng thời tránh từ 'độ chính xác' cho một ảnh đơn.",
    )
    doc.add_paragraph("2.7. Rò rỉ dữ liệu và protocol hợp lệ", style="Heading 2")
    add_body(
        doc,
        "Nếu các frame từ cùng video xuất hiện ở cả train và test, model có thể học danh tính, nền hoặc codec thay vì dấu vết giả mạo. Protocol tốt phải chia theo video hoặc identity trước khi trích frame, giữ test set bất biến và không chỉnh ngưỡng dựa trên test. Trong dự án này, demo-source frames được loại khỏi benchmark; giới hạn tương quan frame vẫn được báo cáo rõ.",
    )


def chapter_3(doc: Document) -> None:
    add_chapter(doc, 3, "Dữ liệu và nguồn gốc model")
    doc.add_paragraph("3.1. Dữ liệu trong dự án", style="Heading 2")
    add_body(
        doc,
        "Tập DeepFake Facial có hai thư mục Fake và Real, mỗi thư mục 1.192 frame. Celeb-DF-v2 sample có 100 fake và 100 real. Bộ demo có 18 ảnh trong một thư mục; tên số lẻ là real và số chẵn là fake. Tám frame fake đã sao chép sang demo bị loại khỏi cross-dataset benchmark bằng danh sách tên tệp cố định.",
    )
    add_table(
        doc,
        ["Tập", "Đường dẫn", "Real", "Fake", "Vai trò"],
        [
            ["DeepFake Facial", "data/benchmarks/deepfake-facial/Deep_Fakes", "1.192", "1.192", "Benchmark ngoài demo"],
            ["Celeb-DF-v2 sample", "data/benchmarks/celebdf-v2-*-sample", "100", "100", "Cross-domain"],
            ["Demo images", "data/demo-images", "9", "9", "Trình diễn pipeline"],
        ],
        [1700, 3260, 900, 900, 2600],
        numeric={2, 3},
    )
    add_callout(
        doc,
        "Khi giảng viên hỏi dataset",
        "Mở data/benchmarks để chỉ tập đánh giá; mở data/demo-images/SOURCES.md để chỉ nguồn ảnh demo; mở reports/cross_dataset_results.json để chỉ kết quả từng frame.",
        fill=PALE_TEAL,
        border=GREEN,
    )
    add_figure(doc, "07_dataset_size.png", "Hình 3.1. Quy mô frame trong cross-dataset evaluation.")
    doc.add_paragraph("3.2. Checkpoint và phát biểu về training", style="Heading 2")
    add_body(
        doc,
        "Hệ thống có trọng số AI thật và suy luận PyTorch thật. Checkpoint chính nằm tại models/hf/koreapeter-ms-eff-gcvit-b0-ffpp/model.safetensors; checkpoint phụ tại models/deepfake_detector.pt. Server nạp trọng số khi khởi động nên lần đầu chậm hơn, các lần sau dùng model đã warm.",
    )
    add_body(
        doc,
        "Cần trả lời trung thực: model chính là checkpoint công khai do tác giả huấn luyện trên FaceForensics++, không phải model nhóm train từ random initialization. Checkpoint phụ là EfficientNet-B0 chuyển đổi có metadata validation/test. Nhóm thực hiện chọn model, thử nghiệm, tích hợp, hiệu chỉnh ngưỡng và đánh giá. Repository không có full training pipeline tái tạo model chính từ video gốc, nên không nhận công huấn luyện từ đầu.",
    )
    add_table(
        doc,
        ["Thành phần", "Kiến trúc", "Nguồn train", "Lưu tại", "Vai trò"],
        [
            ["Chính", "MS-EffGCViT B0", "FF++ - checkpoint công khai", "models/hf/.../model.safetensors", "Quyết định chính"],
            ["Phụ", "EfficientNet-B0", "FF++ C23 - metadata .pt", "models/deepfake_detector.pt", "Cảnh báo mặt nhỏ"],
        ],
        [1200, 1600, 2150, 2600, 1810],
    )
    doc.add_paragraph("3.3. Metadata checkpoint phụ", style="Heading 2")
    add_body(
        doc,
        "Checkpoint phụ ghi validation 450 mẫu và test 450 mẫu, chia theo video nguồn trước khi tách. Test accuracy 82,22%, AUC 89,44%, precision 85,02%, recall 78,22% và F1 81,48%. Đây là metadata checkpoint, không phải cross-dataset score của full pipeline.",
    )
    add_table(
        doc,
        ["Split", "N", "Accuracy", "AUC", "Precision", "Recall", "F1"],
        [
            ["Validation", "450", "79,11%", "89,46%", "88,30%", "67,11%", "76,26%"],
            ["Test", "450", "82,22%", "89,44%", "85,02%", "78,22%", "81,48%"],
        ],
        [1200, 700, 1400, 1300, 1500, 1300, 1960],
        numeric={1, 2, 3, 4, 5, 6},
    )
    doc.add_paragraph("3.4. Bộ ảnh demo", style="Heading 2")
    add_body(
        doc,
        "Ảnh demo được tuyển chọn để minh họa thao tác nên không đại diện ảnh ngoài Internet. Ảnh thật là chân dung công khai; ảnh fake chẵn chủ yếu là frame thao túng từ DeepFake Facial và một ảnh face-swap do chủ dự án cung cấp. SOURCES.md lưu URL hoặc ánh xạ frame. Quy ước lẻ/chẵn chỉ làm ground truth cho script demo, API không dùng tên file.",
    )
    add_figure(doc, "05_demo_montage.png", "Hình 3.2. Mẫu real và fake trong thư mục demo.")


def chapter_4(doc: Document) -> None:
    add_chapter(doc, 4, "Phân tích và thiết kế hệ thống")
    doc.add_paragraph("4.1. Yêu cầu chức năng", style="Heading 2")
    add_bullets(
        doc,
        [
            "Chọn hoặc kéo thả ảnh, giới hạn 16 MB; từ chối tệp không giải mã được.",
            "Phát hiện mặt và trả uncertain nếu không có crop đủ rõ.",
            "Chạy model chính và phụ cục bộ, không phụ thuộc AI bên ngoài.",
            "Hiển thị kết luận, risk, margin, risk band, face status và scan ID.",
            "Sinh tín hiệu riêng theo raw score, ngưỡng, kích thước mặt và chất lượng ảnh.",
        ],
    )
    doc.add_paragraph("4.2. Yêu cầu phi chức năng", style="Heading 2")
    add_bullets(
        doc,
        [
            "Giao diện một trang, ưu tiên upload - quét - đọc kết quả.",
            "Warm model khi khởi động để giảm độ trễ các lần sau.",
            "Cùng ảnh và cùng model phải cho kết quả xác định.",
            "Không lưu ảnh upload lâu dài; xóa tệp tạm sau suy luận.",
            "Có test API, validation, schema, fallback và benchmark tái chạy.",
        ],
    )
    doc.add_paragraph("4.3. Kiến trúc tổng thể", style="Heading 2")
    add_figure(doc, "01_architecture.png", "Hình 4.1. Kiến trúc triển khai FaceTrust.")
    add_body(
        doc,
        "Frontend HTML/CSS/JavaScript gọi POST /api/detect bằng multipart form-data. FastAPI đọc tối đa 16 MB, xác thực ảnh, ghi thư mục tạm và chuyển suy luận CPU sang worker thread. Detector trả DetectionResult có raw evidence. detector_storage chỉ công khai trường cần cho UI, tạo scan ID từ hash nội dung và không tiết lộ đường dẫn tạm.",
    )
    doc.add_paragraph("4.4. Full pipeline", style="Heading 2")
    add_figure(doc, "02_full_pipeline.png", "Hình 4.2. Full pipeline từ ảnh đầu vào đến kết luận.")
    add_numbered(
        doc,
        [
            "Nhận ảnh, kiểm tra dung lượng và phần mở rộng.",
            "Giải mã ảnh thật để chặn tệp giả mạo định dạng.",
            "Phát hiện khuôn mặt và chọn vùng mặt chính.",
            "Crop có margin, resize 224×224 và chuẩn hóa tensor.",
            "Chạy detector chính và checkpoint phụ.",
            "Áp dụng ngưỡng cùng quy tắc khuôn mặt nhỏ.",
            "Tạo score, tín hiệu giải thích và phản hồi JSON.",
        ],
    )
    doc.add_paragraph("4.5. Quy tắc kết luận", style="Heading 2")
    add_body(
        doc,
        "Gọi p là raw fake score detector chính; a là score fake detector phụ; r là tỷ lệ diện tích khuôn mặt trên diện tích ảnh. Ngưỡng hiện tại: t_p = 0,50; t_a = 0,30; t_r = 0,12.",
    )
    add_code_block(
        doc,
        "FAKE nếu có mặt và [p >= 0,50 hoặc (p < 0,50 và r < 0,12 và a >= 0,30)]\n"
        "REAL nếu có mặt và không thỏa điều kiện FAKE\n"
        "UNCERTAIN nếu không định vị được khuôn mặt",
    )
    add_body(
        doc,
        "Detector phụ không tự quyết định trên khuôn mặt đủ lớn. Quy tắc được thêm sau khi kiểm tra face-swap có khuôn mặt nhỏ: model chính dễ bị nền và trang phục chi phối. Điều kiện r < 12% giới hạn ảnh hưởng nhánh phụ để tránh tăng false positive trên chân dung rõ.",
    )
    doc.add_paragraph("4.6. Fake risk index", style="Heading 2")
    add_code_block(
        doc,
        "R_p(p) = p, nếu p < 0,50\n"
        "R_p(p) = 0,50 + 0,30 × (p - 0,50) / 0,50, nếu p >= 0,50\n"
        "R_a(a) = 0,50 + 0,20 × (a - 0,30) / 0,70, nếu cảnh báo mặt nhỏ\n"
        "Risk = max(R_p, R_a)",
    )
    add_figure(doc, "03_score_mapping.png", "Hình 4.3. Ánh xạ raw score sang fake risk index.")
    add_body(
        doc,
        "Nhánh chính đạt tối đa 0,80 và nhánh phụ tối đa 0,70. Khoảng hiển thị được nén vì raw score chưa calibration. UI không dùng 99-100% như mức chắc chắn tuyệt đối. Risk band chỉ là quy ước trình bày, không phải xác suất.",
    )
    doc.add_paragraph("4.7. Decision margin", style="Heading 2")
    add_code_block(
        doc,
        "M = (p - 0,50) / 0,50, nếu nhánh chính kết luận fake\n"
        "M = (a - 0,30) / 0,70, nếu cảnh báo mặt nhỏ kết luận fake\n"
        "M = (0,50 - p) / 0,50, nếu kết luận real\n"
        "M = 0, nếu uncertain",
    )
    add_callout(
        doc,
        "Cách đọc",
        "M = 0 nghĩa score sát ngưỡng; M lớn hơn nghĩa xa ngưỡng theo nhánh đang kích hoạt. M không đo xác suất đúng và không thay thế calibration curve.",
        fill=PALE_BLUE,
        border=BLUE,
    )
    doc.add_paragraph("4.8. Ví dụ tính điểm", style="Heading 2")
    add_body(
        doc,
        "Với ảnh face-swap do chủ dự án cung cấp, một lần đo sau sửa pipeline cho p = 0,6465; a = 0,4145; r = 0,059. Vì p >= 0,50 nên nhãn fake. Risk = 0,50 + 0,30 × (0,6465 - 0,50) / 0,50 = 0,5879, hiển thị 58,8/100. Margin = (0,6465 - 0,50) / 0,50 = 0,293. Đây là ví dụ số học, không phải 58,8% xác suất ảnh fake.",
    )
    doc.add_paragraph("4.9. Tín hiệu giải thích", style="Heading 2")
    add_body(
        doc,
        "Giải thích sinh từ evidence từng ảnh, không dùng câu cố định: raw score và ngưỡng detector chính; score nhánh phụ cùng điều kiện kích hoạt; độ đồng thuận hai nhánh; kích thước crop và tỷ lệ mặt; Laplacian, blockiness, ELA p95 và noise residual. Forensic features hỗ trợ mô tả chất lượng, không tự đảo nhãn model.",
    )
    doc.add_paragraph("4.10. Mode lỗi và phản hồi hệ thống", style="Heading 2")
    add_table(
        doc,
        ["Tình huống", "Kết quả", "Lý do thiết kế"],
        [
            ["Tệp vượt 16 MB", "HTTP 400", "Chặn tải và xử lý không cần thiết"],
            ["Phần mở rộng hợp lệ nhưng byte lỗi", "HTTP 400", "Xác thực bằng giải mã ảnh thật"],
            ["Không phát hiện khuôn mặt", "uncertain", "Không ép classifier đoán ngoài điều kiện đầu vào"],
            ["Hai nhánh chưa đồng thuận", "Giữ quy tắc detector chính", "Detector phụ chỉ có quyền trong chế độ mặt nhỏ"],
        ],
        [2300, 1900, 5160],
    )
    add_body(
        doc,
        "Thiết kế lỗi theo fail-closed ở bước đọc tệp và fail-explicit ở bước nhận diện: tệp không hợp lệ bị từ chối, còn ảnh thiếu khuôn mặt trả uncertain. Cách này giúp phân biệt lỗi hệ thống với giới hạn bằng chứng của ảnh.",
    )


def chapter_5(doc: Document) -> None:
    add_chapter(doc, 5, "Triển khai hệ thống")
    doc.add_paragraph("5.1. Công nghệ", style="Heading 2")
    add_table(
        doc,
        ["Lớp", "Công nghệ", "Vai trò"],
        [
            ["Frontend", "HTML5, CSS3, JavaScript", "Upload, trạng thái quét, kết quả"],
            ["Backend", "Python, FastAPI, Uvicorn", "API, validation, điều phối"],
            ["AI", "PyTorch, torchvision, DeepGuard", "Nạp checkpoint và suy luận"],
            ["Xử lý ảnh", "OpenCV, Pillow, NumPy", "Đọc ảnh, phát hiện mặt, evidence"],
            ["Kiểm thử", "pytest, TestClient, ruff", "Hồi quy API và chất lượng mã"],
        ],
        [1500, 2900, 4960],
    )
    doc.add_paragraph("5.2. Cấu trúc mã nguồn", style="Heading 2")
    add_code_block(
        doc,
        "src/facetrust_benchmark/\n"
        "  web.py                  FastAPI app và endpoint\n"
        "  detector_storage.py     upload pipeline và presentation schema\n"
        "  deepfake_detector.py    model, face crop, score và evidence\n"
        "  image_io.py             validation/đọc ảnh\n"
        "  static/                 HTML, CSS, JavaScript\n"
        "models/                   checkpoint cục bộ\n"
        "data/                     demo và benchmark frames\n"
        "scripts/                  evaluation và tạo báo cáo\n"
        "reports/                  Markdown/JSON kết quả\n"
        "tests/                    kiểm thử tự động",
    )
    doc.add_paragraph("5.3. API và quản lý tệp", style="Heading 2")
    add_body(
        doc,
        "GET /api/health trả trạng thái dịch vụ. POST /api/detect nhận trường image. Backend đọc byte, kiểm tra dung lượng, giải mã bằng image_io, suy luận trong TemporaryDirectory và xóa tệp khi hoàn tất. JSON công khai không chứa raw đường dẫn hoặc note nội bộ nhưng đủ trường để UI giải thích.",
    )
    add_code_block(
        doc,
        '{\n  "scan_id": "6970178998",\n  "label": "fake",\n  "face_detected": true,\n  "presentation": {\n    "fake_risk_index": 0.588,\n    "decision_margin": 0.293,\n    "risk_band": "Cần kiểm tra",\n    "signals": [...]\n  }\n}',
    )
    doc.add_paragraph("5.4. Giao diện", style="Heading 2")
    add_figure(doc, "06_ui_workflow.png", "Hình 5.1. Trạng thái quét và màn hình kết quả.")
    add_body(
        doc,
        "Giao diện được rút về một workflow: chọn ảnh, xem preview, bấm bắt đầu, chờ trạng thái quét và đọc kết luận. Hai vùng đầu vào/kết quả đặt cạnh nhau trên desktop và xếp dọc ở màn hình hẹp. Màu không phải tín hiệu duy nhất: mỗi verdict có ký hiệu chữ, tiêu đề và mô tả.",
    )
    doc.add_paragraph("5.5. Warm-up và hiệu năng", style="Heading 2")
    add_body(
        doc,
        "Ứng dụng gọi warm_detector trong lifespan startup. Nạp safetensors, khởi tạo DeepGuard và checkpoint phụ có thể kéo dài ở lần mở server. Sau khi model ở RAM, mỗi request chỉ thực hiện face detection, transform và hai lần inference. Khi đo thời gian phải tách cold start khỏi warm inference.",
    )
    doc.add_paragraph("5.6. Kiểm thử", style="Heading 2")
    add_bullets(
        doc,
        [
            "Kiểm thử endpoint health và trang HTML.",
            "Kiểm thử ảnh hợp lệ, tệp không phải ảnh và giới hạn dung lượng.",
            "Kiểm thử schema public không lộ raw nội bộ.",
            "Kiểm thử detector trả tập nhãn cho phép và evidence cần thiết.",
            "Chạy ruff; chạy pytest với plugin autoload bị tắt.",
        ],
    )


def chapter_6(doc: Document, cross: dict, pipeline: dict) -> None:
    add_chapter(doc, 6, "Evaluation, benchmark và kết quả")
    doc.add_paragraph("6.1. Ba lớp đánh giá", style="Heading 2")
    add_body(
        doc,
        "Evaluation là quá trình đo theo protocol xác định: dữ liệu, ground truth, ngưỡng và metric. Benchmark là bộ dữ liệu cộng protocol dùng để so sánh lặp lại. Full-pipeline evaluation kiểm tra cả upload, validation, face detection, model, ánh xạ kết quả và API, không chỉ forward pass.",
    )
    add_table(
        doc,
        ["Lớp", "Đo điều gì", "Bằng chứng"],
        [
            ["Model card", "Số liệu tác giả checkpoint", "models/hf/.../README.md"],
            ["Cross-dataset", "Khái quát hóa ngoài demo", "reports/cross_dataset_results.json"],
            ["Full pipeline", "API, face detection, timing", "reports/evaluation_results.json"],
            ["Test", "Hành vi mã và hồi quy", "tests/"],
        ],
        [1800, 3900, 3660],
    )
    doc.add_paragraph("6.2. Protocol cross-dataset", style="Heading 2")
    add_numbered(
        doc,
        [
            "Nạp model giống web bằng warm_detector.",
            "Đọc toàn bộ frame hợp lệ ở hai cặp thư mục fake/real.",
            "Loại tám frame fake đã sao chép sang demo bằng tên chính xác.",
            "Lưu label, risk, margin, face status và latency từng frame.",
            "Tính strict accuracy với uncertain là sai; tính recall, F1 và uncertain riêng.",
            "Ghi JSON chi tiết và Markdown có thời điểm, backend, protocol.",
        ],
    )
    add_figure(doc, "07_dataset_size.png", "Hình 6.1. Số frame dùng cho cross-dataset evaluation.")
    add_callout(
        doc,
        "Giới hạn",
        "Đây là frame-level evaluation. Nhiều frame có thể thuộc cùng video. Benchmark mạnh hơn cần split theo video/identity và giữ test set không tham gia hiệu chỉnh ngưỡng.",
    )
    doc.add_paragraph("6.3. Không trộn số liệu", style="Heading 2")
    add_table(
        doc,
        ["Nguồn", "Dữ liệu", "Accuracy", "AUC", "Diễn giải"],
        [
            ["Tác giả B0", "FF++ cùng miền", "98,08%", "99,69%", "Không phải kết quả nhóm"],
            ["Tác giả B0", "Celeb-DF-v2 cross", "72,59%", "69,99%", "Model card"],
            ["Checkpoint phụ", "FF++ test metadata", "82,22%", "89,44%", "Metadata .pt"],
            [
                "FaceTrust",
                "Hai tập local cross",
                pct(cross.get("overall", {}).get("strict_accuracy", 0)),
                "Không tính",
                "Kết quả nhóm tái chạy",
            ],
        ],
        [1900, 2100, 1200, 1100, 3060],
        numeric={2, 3},
    )
    doc.add_paragraph("6.4. Kết quả cross-dataset", style="Heading 2")
    rows = []
    for result in cross.get("datasets", []):
        summary = result["summary"]
        rows.append(
            [
                result["name"],
                summary["total"],
                pct(summary["strict_accuracy"]),
                pct(summary["balanced_accuracy"]),
                pct(summary["fake_recall"]),
                pct(summary["real_recall"]),
                pct(summary["uncertain_rate"]),
            ]
        )
    overall = cross.get("overall", {})
    if overall:
        rows.append(
            [
                "Gộp",
                overall["total"],
                pct(overall["strict_accuracy"]),
                pct(overall["balanced_accuracy"]),
                pct(overall["fake_recall"]),
                pct(overall["real_recall"]),
                pct(overall["uncertain_rate"]),
            ]
        )
    add_table(
        doc,
        ["Dataset", "N", "Strict Acc.", "Balanced", "Fake recall", "Real recall", "Uncertain"],
        rows,
        [1800, 700, 1300, 1450, 1350, 1350, 1410],
        numeric={1, 2, 3, 4, 5, 6},
    )
    add_figure(doc, "08_cross_dataset_metrics.png", "Hình 6.2. Accuracy và recall theo dataset.")
    add_body(
        doc,
        "Kết quả phải đọc theo từng lớp. Fake recall cao hơn real recall nghĩa detector ưu tiên bắt ảnh fake nhưng đánh nhầm một phần ảnh thật. Balanced accuracy lấy trung bình recall hai lớp nên phản ánh cân bằng tốt hơn một accuracy duy nhất.",
    )
    doc.add_paragraph("6.5. Ma trận nhầm lẫn", style="Heading 2")
    if overall:
        cm = overall["confusion_matrix"]
        add_table(
            doc,
            ["Ground truth", "Predicted fake", "Predicted real", "Uncertain"],
            [
                ["Fake", cm["true_fake_pred_fake"], cm["true_fake_pred_real"], cm["true_fake_pred_uncertain"]],
                ["Real", cm["true_real_pred_fake"], cm["true_real_pred_real"], cm["true_real_pred_uncertain"]],
            ],
            [2160, 2400, 2400, 2400],
            numeric={1, 2, 3},
        )
    add_figure(doc, "09_confusion_matrix.png", "Hình 6.3. Ma trận nhầm lẫn gộp.")
    add_body(
        doc,
        "True fake/Predicted fake là ảnh thao túng bắt đúng. True real/Predicted fake là false positive. Uncertain được tính sai trong strict accuracy. Khi sàng lọc deepfake cần giảm false negative, nhưng false positive cũng gây hậu quả nên không thể bỏ qua.",
    )
    doc.add_paragraph("6.6. Full-pipeline evaluation", style="Heading 2")
    summary = pipeline.get("summary", {})
    if summary:
        latency = summary.get("latency_ms", {})
        add_table(
            doc,
            ["Chỉ số", "Giá trị", "Ý nghĩa"],
            [
                ["Số ảnh demo", summary.get("total", 0), "Trình diễn, không công bố accuracy tổng quát"],
                ["API success", pct(summary.get("api_success_rate", 0)), "POST trả HTTP 200"],
                ["Face detection", pct(summary.get("face_detection_rate", 0)), "Có crop khuôn mặt"],
                ["Median latency", f"{latency.get('median', 0):.1f} ms", "Request sau warm-up"],
                ["P95 latency", f"{latency.get('p95', 0):.1f} ms", "Phân vị 95"],
            ],
            [2100, 1700, 5560],
            numeric={1},
        )
    add_body(
        doc,
        "Demo tuyển chọn chứng minh hệ thống hoạt động end-to-end. Bằng chứng khái quát hóa lấy từ cross-dataset. Hai phép đánh giá có mục tiêu khác nhau và không được trộn thành một accuracy đẹp mắt.",
    )
    doc.add_paragraph("6.7. Phân tích sai số", style="Heading 2")
    add_bullets(
        doc,
        [
            "Ảnh thật có texture/nén gần artifact đã học có thể bị cảnh báo fake.",
            "Ảnh fake từ công cụ mới hoặc hậu xử lý mạnh có thể bị bỏ sót.",
            "Khuôn mặt nhỏ làm tín hiệu vùng mặt yếu; detector phụ chỉ giảm một phần.",
            "Haar cascade có thể bỏ mặt nghiêng, che khuất hoặc quá nhỏ.",
            "Frame cùng video tương quan làm số mẫu lớn hơn số tình huống độc lập.",
        ],
    )
    add_callout(
        doc,
        "Đánh giá mức độ model",
        "Phiên bản hiện tại phù hợp đồ án và sàng lọc có người kiểm tra lại; chưa đủ dùng độc lập trong pháp lý. Ưu điểm là pipeline tái lập và fake recall đáng kể; nhược điểm là real recall/cross-domain chưa ổn định.",
        fill="FBECE9",
        border=RED,
    )
    doc.add_paragraph("6.8. Tính tái lập", style="Heading 2")
    add_body(
        doc,
        "Script ghi thời điểm tạo, tên backend, protocol lấy mẫu, thư mục dữ liệu, danh sách tệp bị loại và toàn bộ row dự đoán. Vì vậy người chấm có thể tính lại summary từ JSON hoặc chạy lại script. Khi công bố con số, cần gắn nó với commit chứa checkpoint, ngưỡng và report tương ứng.",
    )
    doc.add_paragraph("6.9. Đe dọa tính hợp lệ", style="Heading 2")
    add_table(
        doc,
        ["Nhóm rủi ro", "Biểu hiện", "Cách giảm"],
        [
            ["Selection bias", "Demo đã tuyển chọn", "Không dùng demo cho accuracy"],
            ["Frame correlation", "Nhiều frame cùng video", "Báo cáo giới hạn; tương lai split theo video"],
            ["Domain mismatch", "Train FF++, test nguồn khác", "Báo cáo cross-dataset riêng"],
            ["Threshold overfit", "Chỉnh ngưỡng theo test", "Giữ calibration set độc lập"],
        ],
        [2100, 3100, 4160],
    )


def chapter_7(doc: Document) -> None:
    add_chapter(doc, 7, "Quy trình làm việc với AI")
    doc.add_paragraph("7.1. Vai trò AI hỗ trợ", style="Heading 2")
    add_body(
        doc,
        "AI được dùng như trợ lý kỹ thuật qua nhiều phiên: phân tích yêu cầu, khảo sát model/dataset, đọc mã, đề xuất kiến trúc, viết và sửa mã, tạo test, phân tích kết quả và soạn tài liệu. AI không phải ground truth. Mọi thay đổi cuối cùng được đối chiếu file trên máy, chạy lệnh kiểm thử và xem output thật.",
    )
    add_figure(doc, "04_ai_workflow.png", "Hình 7.1. Quy trình phát triển có AI hỗ trợ.")
    doc.add_paragraph("7.2. Các giai đoạn", style="Heading 2")
    add_numbered(
        doc,
        [
            "Làm rõ yêu cầu: chuyển từ bảo vệ ảnh sang nhận diện ảnh sau deepfake.",
            "Khảo sát: thử checkpoint, tải dữ liệu, kiểm tra CPU và ảnh ngoài demo.",
            "Thiết kế lại: rút UI về upload - scan - verdict; tách demo khỏi benchmark; bỏ Gemini khi không tạo giá trị ổn định.",
            "Triển khai: face crop theo model card, detector phụ có điều kiện và ngưỡng công khai.",
            "Đánh giá: test, demo pipeline, cross-dataset và phân tích FP/FN.",
            "Tài liệu hóa: README, đặc tả, quy trình AI, báo cáo JSON/Markdown và Word.",
        ],
    )
    doc.add_paragraph("7.3. Vòng lặp phát hiện và sửa lỗi", style="Heading 2")
    add_body(
        doc,
        "Một ảnh face-swap có khuôn mặt nhỏ từng bị kết luận real. AI hỗ trợ đọc evidence và nhận ra model chính đang bị nền chi phối. Nhóm chuyển sang crop 224×224 theo model card, sau đó thêm checkpoint phụ khi r < 12% và a >= 0,30. Ảnh mẫu được sửa, nhưng nhóm không dừng ở một ca mà chạy cross-dataset để đo cả fake recall và real recall.",
    )
    doc.add_paragraph("7.4. Kiểm soát chất lượng", style="Heading 2")
    add_bullets(
        doc,
        [
            "Không chấp nhận claim của AI nếu thiếu file, dòng mã hoặc output lệnh.",
            "Không đưa API key vào repository; bỏ dịch vụ ngoài không ổn định.",
            "Không dùng tên tệp real/fake trong API để quyết định nhãn.",
            "Không tối ưu bộ demo rồi gọi là benchmark.",
            "Giữ test hồi quy và ruff trước khi chốt.",
            "Ghi giới hạn và trường hợp thất bại, không chỉ chụp kết quả đẹp.",
        ],
    )
    doc.add_paragraph("7.5. Giá trị và giới hạn", style="Heading 2")
    add_body(
        doc,
        "AI tăng tốc khảo sát và triển khai, nhất là khi yêu cầu thay đổi nhiều lần. Tuy nhiên AI có thể lẫn số liệu model card với kết quả dự án hoặc tạo giao diện thuyết phục hơn chất lượng model. Cơ chế phòng ngừa là tách bằng chứng, tái chạy script và để mã nguồn quyết định thay vì lời mô tả.",
    )
    doc.add_paragraph("7.6. Nhật ký quyết định kỹ thuật", style="Heading 2")
    add_table(
        doc,
        ["Quyết định", "Bằng chứng", "Kết quả"],
        [
            ["Bỏ Gemini khỏi pipeline", "Không cải thiện ổn định, tăng độ trễ và phụ thuộc key", "Hai detector cục bộ"],
            ["Dùng face crop", "Model card yêu cầu face-cropped 224×224", "Giảm ảnh hưởng nền"],
            ["Tách demo/benchmark", "Demo đã chọn theo khả năng trình diễn", "Không dùng demo cho claim tổng quát"],
            ["Đổi confidence thành margin", "Score chưa calibration", "UI bớt tạo cảm giác chắc chắn giả"],
        ],
        [2600, 3820, 2940],
    )
    doc.add_paragraph("7.7. Dấu vết làm việc", style="Heading 2")
    add_body(
        doc,
        "Dấu vết được giữ trong lịch sử mã nguồn, scripts đánh giá, report JSON/Markdown, tài liệu AI_WORKFLOW.md và kiểm thử. Báo cáo không chép nguyên hội thoại dài; thay vào đó tóm tắt quyết định, nguyên nhân, phép đo và thay đổi có thể kiểm chứng trong repository.",
    )


def chapter_8(doc: Document, cross: dict) -> None:
    add_chapter(doc, 8, "Kết luận và hướng phát triển")
    overall = cross.get("overall", {})
    result = (
        f" Trên {overall['total']} frame cross-dataset, strict accuracy đạt {pct(overall['strict_accuracy'])}, fake recall {pct(overall['fake_recall'])} và real recall {pct(overall['real_recall'])}."
        if overall
        else ""
    )
    doc.add_paragraph("8.1. Kết luận", style="Heading 2")
    add_body(
        doc,
        "Đồ án đã xây dựng full pipeline gồm upload, validation, phát hiện mặt, crop, hai checkpoint, kết luận ba trạng thái, score có công thức, evidence, API và giao diện." + result,
    )
    add_body(
        doc,
        "Anti-deepfake không bất khả thi nhưng rất khó đạt tin cậy cao ngoài miền huấn luyện. Same-dataset score cao không đảm bảo ảnh Internet. Giá trị lớn của phiên bản hiện tại là tính minh bạch: model thật, trọng số cục bộ, dữ liệu có đường dẫn, script tái chạy, ma trận nhầm lẫn và giới hạn rõ.",
    )
    doc.add_paragraph("8.2. Hạn chế", style="Heading 2")
    add_bullets(
        doc,
        [
            "Model chính là pretrained checkpoint, chưa fine-tune lại trên tập đa nguồn lớn.",
            "Cross-dataset còn thiên lệch fake recall và real recall.",
            "Benchmark frame-level có tương quan cùng video.",
            "Chưa có calibration set để biến score thành xác suất thống kê.",
            "Chưa xử lý temporal inconsistency của video.",
            "Evidence forensic là mô tả hỗ trợ, chưa phải giải thích nhân quả.",
        ],
    )
    doc.add_paragraph("8.3. Hướng phát triển", style="Heading 2")
    add_numbered(
        doc,
        [
            "Tạo train/validation/test theo video và identity từ nhiều nguồn.",
            "Fine-tune với augmentation nén mạng xã hội, resize, blur, screenshot và re-encoding.",
            "Hiệu chỉnh xác suất bằng temperature scaling hoặc isotonic regression.",
            "Báo cáo ROC/PR curve và chọn threshold theo chi phí FP/FN.",
            "Thêm video pipeline, tracking mặt và tổng hợp score theo thời gian.",
            "Thêm Grad-CAM sau khi kiểm tra độ ổn định, không dùng để trang trí.",
            "Version hóa model, dataset manifest và experiment log.",
        ],
    )


def appendices(doc: Document, cross: dict) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph("PHỤ LỤC A. HƯỚNG DẪN TÁI LẬP", style="Heading 1")
    paragraph_rule(heading)
    doc.add_paragraph("A.1. Chạy web", style="Heading 2")
    add_code_block(doc, r"cd D:\Xampp\htdocs\anti-deepfake-face" + "\n.\\start-web.ps1\n# Mở http://127.0.0.1:8000")
    doc.add_paragraph("A.2. Chạy kiểm thử", style="Heading 2")
    add_code_block(
        doc,
        ".\\.venv\\Scripts\\python.exe -m ruff check src tests scripts\n"
        "$env:PYTEST_DISABLE_PLUGIN_AUTOLOAD='1'\n"
        ".\\.venv\\Scripts\\python.exe -m pytest -q",
    )
    doc.add_paragraph("A.3. Chạy evaluation", style="Heading 2")
    add_code_block(
        doc,
        "$env:PYTHONPATH='D:\\Xampp\\htdocs\\anti-deepfake-face\\src'\n"
        ".\\.venv\\Scripts\\python.exe scripts\\evaluate_system.py\n"
        ".\\.venv\\Scripts\\python.exe scripts\\evaluate_cross_dataset.py",
    )
    add_body(
        doc,
        "Kết quả ở reports/*.md và reports/*.json. JSON chứa summary theo dataset và rows từng frame. Khi đổi checkpoint hoặc ngưỡng, phải chạy lại benchmark và commit kết quả cùng mã tương ứng.",
    )
    doc.add_paragraph("A.4. Môi trường và tệp đầu ra", style="Heading 2")
    add_table(
        doc,
        ["Thành phần", "Vị trí / lệnh kiểm tra", "Mục đích"],
        [
            ["Python virtual env", ".venv/", "Cô lập thư viện chạy web và model"],
            ["Model chính", "models/hf/.../model.safetensors", "Trọng số face-forgery"],
            ["Model phụ", "models/deepfake_detector.pt", "Bổ trợ khuôn mặt nhỏ"],
            ["Cross report", "reports/cross_dataset_results.*", "Metric và row từng frame"],
            ["Pipeline report", "reports/evaluation_results.*", "API, face detection, latency"],
        ],
        [2000, 3760, 3600],
    )

    doc.add_page_break()
    heading = doc.add_paragraph("PHỤ LỤC B. CÂU HỎI PHẢN BIỆN", style="Heading 1")
    paragraph_rule(heading)
    questions = [
        (
            "1. Dự án có dataset và train AI thật không?",
            "Có dataset trong data/benchmarks và trọng số AI trong models. Model chính do tác giả checkpoint train trên FaceForensics++; nhóm không train từ đầu. Nhóm chọn model, tích hợp, hiệu chỉnh và đánh giá. Checkpoint phụ có metadata validation/test trong .pt.",
        ),
        (
            "2. Làm sao biết dùng model thật, không trả đại?",
            "deepfake_detector.py nạp model.safetensors, tạo tensor 224×224 và gọi model(tensor).sigmoid(). Đổi pixel làm đổi score. JSON có score, face status, latency và scan ID. Đổi tên cùng ảnh không đổi nội dung dự đoán.",
        ),
        (
            "3. Làm sao biết đánh giá không cảm tính?",
            "Ground truth đến từ thư mục Real/Fake; script tính TP/FP/TN/FN. Demo bị loại khỏi benchmark. JSON lưu từng dòng để kiểm tra thủ công; Markdown sinh từ cùng dữ liệu.",
        ),
        (
            "4. Mức độ model hiện tại?",
            "Đủ cho đồ án và sàng lọc có người xem lại, chưa đủ pháp chứng độc lập. Cần trình bày recall từng lớp, balanced accuracy và confusion matrix, không chỉ accuracy.",
        ),
        (
            "5. Evaluation và benchmark khác nhau?",
            "Evaluation là quá trình đo; benchmark là dữ liệu cộng protocol lặp lại. Benchmark quy định split, ground truth, preprocessing, threshold và metric. Demo không phải benchmark.",
        ),
        (
            "6. Full pipeline evaluation là gì?",
            "Kiểm tra đường đi upload đến JSON/UI: validation, face detection, model, score mapping, API success và latency. Model evaluation chỉ đo phân loại.",
        ),
        (
            "7. Vì sao model card 98% nhưng dự án thấp hơn?",
            "98,08% là same-dataset FF++; ảnh ngoài miền khác về người, công cụ giả, nén và crop. Model card cũng công bố cross Celeb-DF-v2 thấp hơn. Đó là domain shift.",
        ),
        (
            "8. Risk 60/100 có nghĩa 60% fake?",
            "Không. Đó là score ánh xạ theo công thức. Muốn gọi xác suất cần calibration set và kiểm tra reliability/Brier score.",
        ),
        (
            "9. Vì sao có detector phụ?",
            "Model chính có thể bỏ sót mặt nhỏ. Detector phụ chỉ kích hoạt khi mặt dưới 12% và score trên 0,30, không chi phối mọi ảnh.",
        ),
        (
            "10. Muốn nâng độ chính xác cần gì?",
            "Đa dạng nguồn train, split theo video/identity, augmentation giống Internet, calibration/test độc lập, tối ưu threshold và temporal aggregation cho video.",
        ),
        (
            "11. Vì sao không dùng bộ demo để tính accuracy?",
            "Demo đã tuyển chọn để trình diễn nên selection bias rất lớn. Accuracy trên demo chỉ cho biết pipeline xử lý được các ca đã chọn, không đo khả năng tổng quát với ảnh bất kỳ.",
        ),
        (
            "12. Uncertain được tính thế nào?",
            "Uncertain xuất hiện khi không định vị được khuôn mặt. Trong strict accuracy nó được tính sai; đồng thời uncertain rate được báo cáo riêng để không giấu giới hạn đầu vào.",
        ),
        (
            "13. Tại sao balanced accuracy bằng trung bình hai recall?",
            "Mỗi lớp đóng góp một nửa bất kể số lượng mẫu. Chỉ số này giúp phát hiện model thiên về fake hoặc real dù accuracy tổng có vẻ chấp nhận được.",
        ),
        (
            "14. Có thể chứng minh tên tệp không quyết định kết quả không?",
            "Sao chép cùng ảnh sang tên khác rồi gọi API. scan ID được tạo từ nội dung byte và detector đọc pixel; quy ước lẻ/chẵn chỉ tồn tại trong script đánh giá demo.",
        ),
        (
            "15. Vì sao lần quét đầu chậm?",
            "Cold start phải nạp hai checkpoint và khởi tạo kiến trúc. Server warm model trong lifespan; latency thường xuyên nên đo sau bước này và báo cold-start riêng.",
        ),
    ]
    for question, answer in questions:
        doc.add_paragraph(question, style="Heading 2")
        add_body(doc, answer)

    doc.add_page_break()
    heading = doc.add_paragraph("PHỤ LỤC C. CÔNG THỨC VÀ VÍ DỤ", style="Heading 1")
    paragraph_rule(heading)
    add_table(
        doc,
        ["Ký hiệu", "Giá trị / công thức", "Ý nghĩa"],
        [
            ["p", "Sigmoid detector chính", "Raw fake score trên face crop"],
            ["a", "Output detector phụ", "Tín hiệu phụ trợ"],
            ["r", "Diện tích mặt / ảnh", "Kiểm tra mặt nhỏ"],
            ["t_p", "0,50", "Ngưỡng chính"],
            ["t_a", "0,30", "Ngưỡng phụ khi mặt nhỏ"],
            ["t_r", "0,12", "Ngưỡng tỷ lệ mặt"],
            ["Risk", "max(R_p, R_a)", "Chỉ số rủi ro"],
            ["Margin", "Khoảng cách tới ngưỡng", "Không phải xác suất"],
        ],
        [1100, 3260, 5000],
    )
    overall = cross.get("overall", {})
    if overall:
        add_callout(
            doc,
            "Bộ số liệu đang dùng",
            f"N = {overall['total']}; correct = {overall['correct']}; strict accuracy = {overall['correct']} / {overall['total']} = {pct(overall['strict_accuracy'])}; balanced accuracy = (fake recall {pct(overall['fake_recall'])} + real recall {pct(overall['real_recall'])}) / 2 = {pct(overall['balanced_accuracy'])}.",
            fill=PALE_TEAL,
            border=GREEN,
        )
    doc.add_paragraph("C.1. Bốn trường hợp quyết định", style="Heading 2")
    add_table(
        doc,
        ["Trường hợp", "Điều kiện ví dụ", "Kết luận", "Điểm được dùng"],
        [
            ["Detector chính vượt ngưỡng", "p = 0,65; có mặt", "fake", "R_p và margin theo p"],
            ["Cảnh báo mặt nhỏ", "p = 0,40; r = 0,08; a = 0,45", "fake", "max(R_p, R_a), margin theo a"],
            ["Vùng real", "p = 0,22; có mặt; không cảnh báo", "real", "R_p = 0,22; margin = 0,56"],
            ["Không có mặt", "face_detected = false", "uncertain", "margin = 0"],
        ],
        [2100, 3100, 1500, 2660],
    )
    doc.add_paragraph("C.2. Ví dụ từ ma trận nhầm lẫn", style="Heading 2")
    add_body(
        doc,
        "Giả sử TP = 152, FP = 122, TN = 73 và 12 mẫu uncertain. Fake precision = 152 / (152 + 122); fake recall = 152 / tổng true fake; real recall = 73 / tổng true real. Strict accuracy chỉ đếm TP + TN trên toàn bộ N, kể cả uncertain. Cách ghi rõ tử số và mẫu số giúp người chấm kiểm tra bằng máy tính mà không cần tin bảng tổng hợp.",
    )


def references(doc: Document) -> None:
    doc.add_page_break()
    heading = doc.add_paragraph("TÀI LIỆU THAM KHẢO", style="Heading 1")
    paragraph_rule(heading)
    refs = [
        (
            "[1] A. Rössler et al., FaceForensics++: Learning to Detect Manipulated Facial Images, ICCV, 2019.",
            "https://openaccess.thecvf.com/content_ICCV_2019/html/Rossler_FaceForensics_Learning_to_Detect_Manipulated_Facial_Images_ICCV_2019_paper.html",
        ),
        (
            "[2] Y. Li et al., Celeb-DF: A Large-Scale Challenging Dataset for DeepFake Forensics, CVPR, 2020.",
            "https://openaccess.thecvf.com/content_CVPR_2020/html/Li_Celeb-DF_A_Large-Scale_Challenging_Dataset_for_DeepFake_Forensics_CVPR_2020_paper.html",
        ),
        (
            "[3] M. Tan and Q. V. Le, EfficientNet: Rethinking Model Scaling for CNNs, ICML, 2019.",
            "https://proceedings.mlr.press/v97/tan19a.html",
        ),
        (
            "[4] HanMoonSub, DeepGuard source repository and model documentation.",
            "https://github.com/HanMoonSub/DeepGuard",
        ),
        (
            "[5] KoreaPeter, MS-EffGCViT B0 Deepfake Detector model card.",
            "https://huggingface.co/KoreaPeter/ms-eff-gcvit-deepfake-b0-ff-plus-plus",
        ),
        ("[6] FastAPI documentation.", "https://fastapi.tiangolo.com/"),
        ("[7] OpenCV documentation.", "https://docs.opencv.org/"),
        (
            "[8] FaceTrust source, local benchmark reports and reproducibility scripts.",
            "https://github.com/thien-2k5/Xu_Ly_Anh",
        ),
    ]
    for text, url in refs:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(7)
        set_run_font(paragraph.add_run(text + " "), size=10, color=INK)
        add_hyperlink(paragraph, url, url)


def apply_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = "FaceTrust - Hệ thống kiểm định ảnh khuôn mặt và phát hiện deepfake"
    props.subject = "Báo cáo đồ án môn Xử lý ảnh"
    props.author = "Nhóm FaceTrust"
    props.keywords = "deepfake, face-swap, benchmark, image forensics"
    props.comments = "Hoàn thiện từ mã nguồn và benchmark tái chạy."


def main() -> None:
    cross = load_json(CROSS_RESULTS)
    pipeline = load_json(PIPELINE_RESULTS)
    doc = Document()
    configure_document(doc)
    apply_metadata(doc)
    build_cover(doc)
    build_front_matter(doc, cross)
    build_toc(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc, cross, pipeline)
    chapter_7(doc)
    chapter_8(doc, cross)
    appendices(doc, cross)
    references(doc)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
